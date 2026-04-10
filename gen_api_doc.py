#!/usr/bin/env python3
"""
Genera Endpoints_PapeleriaCartagena.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Color palette ───────────────────────────────────────────────────────────
C_DARK_BLUE  = RGBColor(0x1E, 0x3A, 0x5F)
C_MID_BLUE   = RGBColor(0x2D, 0x6A, 0xA0)
C_LIGHT_BLUE = RGBColor(0xDF, 0xEA, 0xF5)
C_GREEN      = RGBColor(0x19, 0x6F, 0x3D)
C_LIGHT_GRN  = RGBColor(0xD5, 0xF5, 0xE3)
C_RED        = RGBColor(0x9B, 0x15, 0x15)
C_LIGHT_RED  = RGBColor(0xFC, 0xE4, 0xE4)
C_ORANGE     = RGBColor(0x9C, 0x4A, 0x00)
C_LIGHT_ORG  = RGBColor(0xFF, 0xEE, 0xD5)
C_PURPLE     = RGBColor(0x5B, 0x27, 0x8F)
C_LIGHT_PRP  = RGBColor(0xEF, 0xE6, 0xFB)
C_GRAY_H     = RGBColor(0x2C, 0x2C, 0x2C)
C_GRAY_L     = RGBColor(0xF2, 0xF2, 0xF2)
C_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

METHOD_COLORS = {
    'GET':    (C_LIGHT_BLUE,  C_MID_BLUE),
    'POST':   (C_LIGHT_GRN,   C_GREEN),
    'PUT':    (C_LIGHT_ORG,   C_ORANGE),
    'PATCH':  (C_LIGHT_ORG,   C_ORANGE),
    'DELETE': (C_LIGHT_RED,   C_RED),
}

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_color = str(rgb)  # RGBColor.__str__ returns hex like 'AABBCC'
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)

def set_cell_borders(cell, color='D0D0D0'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(borders)

def cell_text(cell, text, bold=False, color=None, size=10, align=None, italic=False):
    cell.text = ''
    p   = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt(20)
    run.font.color.rgb = C_DARK_BLUE
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt(14)
    run.font.color.rgb = C_MID_BLUE
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt(11)
    run.font.color.rgb = C_GRAY_H
    return p

def normal(text, bold=False, italic=False, color=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color
    return p

def note(text):
    """Highlighted note box (single-row table)."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    set_cell_bg(cell, C_LIGHT_BLUE)
    set_cell_borders(cell, 'A8C4E0')
    p    = cell.paragraphs[0]
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.2)
    run  = p.add_run('ℹ️  ' + text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = C_MID_BLUE
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def warning(text):
    """Warning box."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    set_cell_bg(cell, C_LIGHT_ORG)
    set_cell_borders(cell, 'F0A000')
    p    = cell.paragraphs[0]
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.2)
    run  = p.add_run('⚠️  ' + text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = C_ORANGE
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def divider():
    p = doc.add_paragraph('─' * 90)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(7)

def endpoint_block(method, path, description, auth_roles,
                   path_params=None, query_params=None,
                   request_body=None, response_ok=None,
                   response_err=None, notes_list=None):
    """Renders a full endpoint card as a table."""
    bg, fg = METHOD_COLORS.get(method, (C_GRAY_L, C_GRAY_H))

    # Header row: METHOD + PATH
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Cm(2.0)
    tbl.columns[1].width = Cm(13.5)

    mth_cell = tbl.cell(0, 0)
    set_cell_bg(mth_cell, fg)
    cell_text(mth_cell, method, bold=True, color=C_WHITE,
              size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    pth_cell = tbl.cell(0, 1)
    set_cell_bg(pth_cell, C_GRAY_H)
    cell_text(pth_cell, path, bold=True, color=C_WHITE, size=10)

    # Body table (details)
    body_rows = []

    # Description
    body_rows.append(('Descripción', description))

    # Auth
    if auth_roles:
        body_rows.append(('Autorización', auth_roles))

    # Path params
    if path_params:
        body_rows.append(('Path params', path_params))

    # Query params
    if query_params:
        body_rows.append(('Query params', query_params))

    # Request body
    if request_body:
        body_rows.append(('Request body\n(JSON)', request_body))

    # Response 200
    if response_ok:
        body_rows.append(('Respuesta 200', response_ok))

    # Errors
    if response_err:
        body_rows.append(('Errores comunes', response_err))

    # Notes
    if notes_list:
        body_rows.append(('Notas', notes_list))

    dtbl = doc.add_table(rows=len(body_rows), cols=2)
    dtbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    dtbl.columns[0].width = Cm(3.5)
    dtbl.columns[1].width = Cm(12.0)

    for i, (label, value) in enumerate(body_rows):
        lbl_cell = dtbl.cell(i, 0)
        val_cell = dtbl.cell(i, 1)
        set_cell_bg(lbl_cell, bg)
        set_cell_borders(lbl_cell, 'CCCCCC')
        set_cell_borders(val_cell, 'CCCCCC')
        cell_text(lbl_cell, label, bold=True, color=fg, size=9)
        # Value cell: plain paragraph
        val_cell.text = ''
        p = val_cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(value)
        run.font.size = Pt(9)
        run.font.color.rgb = C_GRAY_H
        val_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Spacing
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after  = Pt(6)
    sp.paragraph_format.space_before = Pt(0)

# ═══════════════════════════════════════════════════════════════════════════
#  DOCUMENT START
# ═══════════════════════════════════════════════════════════════════════════

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run('PAPELERÍA CARTAGENA')
run.bold = True
run.font.size  = Pt(26)
run.font.color.rgb = C_DARK_BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Especificación de API REST — Backend')
run2.font.size = Pt(14)
run2.font.color.rgb = C_MID_BLUE
run2.bold = True

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Documento instructivo para implementación de backend independiente')
run3.font.size = Pt(10)
run3.italic = True
run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
divider()

# ── 1. CONTEXTO ──────────────────────────────────────────────────────────────
h1('1. Contexto del Sistema')

normal(
    'Papelería Cartagena es un sistema B2B de venta mayorista de artículos de papelería. '
    'El frontend está desarrollado en React 18 + Vite y actualmente usa datos en memoria (mock). '
    'Este documento describe todos los endpoints REST que el backend debe exponer para '
    'reemplazar esa capa de datos. El backend puede implementarse en cualquier stack '
    '(Node/Express, Django, Laravel, FastAPI, etc.).'
)

normal('Roles de usuario del sistema:')
roles = [
    ('admin',            'Super-administrador de Papelería Cartagena. Acceso total.'),
    ('advisor',          'Asesor comercial. Gestiona pedidos asignados.'),
    ('client (supervisor)',     'Usuario de empresa cliente con rol supervisor. Aprueba pedidos de su empresa y puede administrar usuarios y sucursales propias.'),
    ('client (creador_pedidos)','Usuario de empresa cliente que sólo puede crear pedidos (quedan en "Pendiente por aprobar" hasta que el supervisor los aprueba).'),
]
tbl = doc.add_table(rows=len(roles)+1, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl.columns[0].width = Cm(4.5)
tbl.columns[1].width = Cm(11.0)
set_cell_bg(tbl.cell(0,0), C_DARK_BLUE); cell_text(tbl.cell(0,0),'Rol',bold=True,color=C_WHITE,size=10)
set_cell_bg(tbl.cell(0,1), C_DARK_BLUE); cell_text(tbl.cell(0,1),'Descripción',bold=True,color=C_WHITE,size=10)
for i,(r,d) in enumerate(roles,1):
    set_cell_bg(tbl.cell(i,0), C_LIGHT_BLUE)
    cell_text(tbl.cell(i,0), r, bold=True, color=C_MID_BLUE, size=9)
    set_cell_borders(tbl.cell(i,0),'CCCCCC')
    set_cell_borders(tbl.cell(i,1),'CCCCCC')
    cell_text(tbl.cell(i,1), d, size=9)
doc.add_paragraph()

note(
    'Base URL sugerida: https://api.papeleriacartagena.com/api/v1  '
    'Todas las respuestas deben devolver Content-Type: application/json. '
    'Las fechas siguen formato ISO 8601 (YYYY-MM-DD o YYYY-MM-DDTHH:MM:SS).'
)

# ── 2. AUTENTICACIÓN ──────────────────────────────────────────────────────────
h1('2. Autenticación')
normal(
    'El sistema utiliza autenticación basada en JWT (JSON Web Tokens). '
    'El token debe enviarse en todas las peticiones protegidas mediante el header: '
    'Authorization: Bearer <token>. El payload del token debe incluir id, email, role y clientRole.'
)

h2('2.1 Login')
endpoint_block(
    method='POST',
    path='/auth/login',
    description='Autentica un usuario y retorna un token JWT junto con los datos del perfil.',
    auth_roles='Público (sin token)',
    request_body=(
        '{\n'
        '  "email":    "usuario@ejemplo.com",   // requerido\n'
        '  "password": "contraseña"             // requerido\n'
        '}'
    ),
    response_ok=(
        '{\n'
        '  "token": "eyJhbGci...",\n'
        '  "user": {\n'
        '    "id": 3,\n'
        '    "name": "Papelería El Centro",\n'
        '    "email": "supervisor@oblicua.com",\n'
        '    "role": "client",\n'
        '    "clientRole": "supervisor",      // null para admin/advisor\n'
        '    "companyId": 1,                  // solo clientes\n'
        '    "sucursalId": 1,                 // solo clientes\n'
        '    "priceListId": 2,                // solo clientes\n'
        '    "initials": "PC"\n'
        '  }\n'
        '}'
    ),
    response_err=(
        '401 { "error": "Credenciales incorrectas" }\n'
        '422 { "error": "email y password son requeridos" }'
    )
)

h2('2.2 Logout')
endpoint_block(
    method='POST',
    path='/auth/logout',
    description='Invalida el token JWT actual (requiere implementar blacklist o rotación de tokens).',
    auth_roles='Cualquier usuario autenticado',
    response_ok='{ "message": "Sesión cerrada correctamente" }'
)

h2('2.3 Me (perfil actual)')
endpoint_block(
    method='GET',
    path='/auth/me',
    description='Retorna los datos del usuario autenticado. Útil para hidratar el contexto al recargar la app.',
    auth_roles='Cualquier usuario autenticado',
    response_ok=(
        '{ "id": 3, "name": "...", "role": "client", "clientRole": "supervisor", ... }'
    )
)

# ── 3. EMPRESAS ──────────────────────────────────────────────────────────────
h1('3. Empresas (Companies)')
normal(
    'Las empresas representan a los clientes corporativos de Papelería Cartagena. '
    'Cada empresa puede tener múltiples sucursales y múltiples usuarios. '
    'El admin gestiona empresas; el supervisor puede leer la suya propia.'
)

h2('3.1 Listar empresas')
endpoint_block(
    method='GET',
    path='/companies',
    description='Retorna todas las empresas registradas. El admin ve todas; el supervisor sólo ve la suya.',
    auth_roles='admin · client/supervisor (solo la propia)',
    query_params=(
        'active=true|false  — Filtrar por estado activo (opcional)\n'
        'search=texto        — Búsqueda por nombre o NIT (opcional)'
    ),
    response_ok=(
        '[\n'
        '  {\n'
        '    "id": 1,\n'
        '    "name": "Papelería El Centro",\n'
        '    "nit": "900.123.456-7",\n'
        '    "email": "contacto@elcentro.com",\n'
        '    "phone": "601-234-5678",\n'
        '    "address": "Cra 10 # 5-23, Bogotá",\n'
        '    "active": true,\n'
        '    "sucursales": [ ... ]            // array anidado\n'
        '  }\n'
        ']'
    )
)

h2('3.2 Crear empresa')
endpoint_block(
    method='POST',
    path='/companies',
    description='Crea una nueva empresa cliente.',
    auth_roles='admin',
    request_body=(
        '{\n'
        '  "name":    "Nombre Empresa",        // requerido\n'
        '  "nit":     "900.000.000-0",         // requerido, único\n'
        '  "email":   "contacto@empresa.com",  // requerido\n'
        '  "phone":   "601-000-0000",          // opcional\n'
        '  "address": "Dirección completa",    // opcional\n'
        '  "active":  true                     // default: true\n'
        '}'
    ),
    response_ok='201 { "id": 3, "name": "...", ... }',
    response_err='409 { "error": "NIT ya registrado" }'
)

h2('3.3 Obtener empresa por ID')
endpoint_block(
    method='GET',
    path='/companies/:id',
    description='Retorna una empresa con sus sucursales anidadas.',
    auth_roles='admin · client/supervisor (solo la propia)',
    path_params=':id — ID entero de la empresa',
    response_ok='{ "id": 1, "name": "...", "sucursales": [ { "id": 1, ... }, ... ] }',
    response_err='404 { "error": "Empresa no encontrada" }'
)

h2('3.4 Editar empresa')
endpoint_block(
    method='PUT',
    path='/companies/:id',
    description='Actualiza los datos de una empresa. Enviar solo los campos a modificar (partial update).',
    auth_roles='admin',
    path_params=':id — ID entero de la empresa',
    request_body=(
        '{\n'
        '  "name":    "Nuevo nombre",  // opcional\n'
        '  "email":   "nuevo@mail.com",\n'
        '  "phone":   "...",\n'
        '  "address": "...",\n'
        '  "active":  false\n'
        '}'
    ),
    response_ok='{ "id": 1, "name": "Nuevo nombre", ... }'
)

h2('3.5 Eliminar empresa')
endpoint_block(
    method='DELETE',
    path='/companies/:id',
    description='Desactiva (soft-delete) o elimina definitivamente una empresa. Se recomienda soft-delete para preservar historial de pedidos.',
    auth_roles='admin',
    path_params=':id — ID entero de la empresa',
    response_ok='{ "message": "Empresa eliminada correctamente" }',
    response_err='409 { "error": "No se puede eliminar: la empresa tiene usuarios activos" }'
)

# ── 4. SUCURSALES ──────────────────────────────────────────────────────────────
h1('4. Sucursales de Empresa')
normal(
    'Las sucursales son sub-recursos de una empresa. Cada usuario tipo cliente '
    'está asociado a una sucursal específica. El supervisor puede gestionar las '
    'sucursales de su propia empresa desde el panel de cliente.'
)

h2('4.1 Listar sucursales de una empresa')
endpoint_block(
    method='GET',
    path='/companies/:companyId/sucursales',
    description='Retorna todas las sucursales de la empresa indicada.',
    auth_roles='admin · client/supervisor (solo la propia empresa)',
    path_params=':companyId — ID entero de la empresa',
    response_ok=(
        '[\n'
        '  {\n'
        '    "id": 1,\n'
        '    "companyId": 1,\n'
        '    "name": "Sede Principal",\n'
        '    "address": "Cra 10 # 5-23, Bogotá",\n'
        '    "city": "Bogotá",\n'
        '    "active": true\n'
        '  }\n'
        ']'
    )
)

h2('4.2 Crear sucursal')
endpoint_block(
    method='POST',
    path='/companies/:companyId/sucursales',
    description='Agrega una nueva sucursal a la empresa.',
    auth_roles='admin · client/supervisor (solo la propia empresa)',
    path_params=':companyId — ID entero de la empresa',
    request_body=(
        '{\n'
        '  "name":    "Sucursal Chapinero",  // requerido\n'
        '  "city":    "Bogotá",              // requerido\n'
        '  "address": "Cra 13 # 56-12",     // opcional\n'
        '  "active":  true                   // default: true\n'
        '}'
    ),
    response_ok='201 { "id": 3, "companyId": 1, "name": "Sucursal Chapinero", ... }'
)

h2('4.3 Editar sucursal')
endpoint_block(
    method='PUT',
    path='/companies/:companyId/sucursales/:sucursalId',
    description='Actualiza los datos de una sucursal.',
    auth_roles='admin · client/supervisor (solo la propia empresa)',
    path_params=':companyId, :sucursalId — IDs enteros',
    request_body='{ "name": "...", "city": "...", "address": "...", "active": true }',
    response_ok='{ "id": 2, "companyId": 1, "name": "...", ... }'
)

h2('4.4 Eliminar sucursal')
endpoint_block(
    method='DELETE',
    path='/companies/:companyId/sucursales/:sucursalId',
    description='Elimina o desactiva una sucursal.',
    auth_roles='admin · client/supervisor (solo la propia empresa)',
    path_params=':companyId, :sucursalId — IDs enteros',
    response_ok='{ "message": "Sucursal eliminada correctamente" }',
    response_err='409 { "error": "No se puede eliminar: la sucursal tiene usuarios asociados" }'
)

# ── 5. USUARIOS ──────────────────────────────────────────────────────────────
h1('5. Usuarios')
normal(
    'Existen tres roles principales: admin, advisor y client. Los usuarios tipo client '
    'tienen además un campo clientRole: "supervisor" o "creador_pedidos". '
    'El admin gestiona todos los usuarios; el supervisor solo gestiona usuarios de su empresa.'
)

warning(
    'NUNCA retornar el campo "password" en ningún endpoint de lectura. '
    'Guardar contraseñas con bcrypt o argon2 (mínimo 12 rondas).'
)

h2('5.1 Listar usuarios')
endpoint_block(
    method='GET',
    path='/users',
    description='Lista usuarios. El admin ve todos; el supervisor ve solo los de su empresa.',
    auth_roles='admin · client/supervisor (filtra por companyId automáticamente)',
    query_params=(
        'role=admin|advisor|client  — Filtrar por rol\n'
        'companyId=1                — Filtrar por empresa (solo admin)\n'
        'sucursalId=1               — Filtrar por sucursal\n'
        'active=true|false          — Filtrar por estado\n'
        'search=texto               — Búsqueda por nombre o email'
    ),
    response_ok=(
        '[\n'
        '  {\n'
        '    "id": 3,\n'
        '    "name": "Papelería El Centro",\n'
        '    "email": "supervisor@oblicua.com",\n'
        '    "role": "client",\n'
        '    "clientRole": "supervisor",\n'
        '    "companyId": 1,\n'
        '    "sucursalId": 1,\n'
        '    "priceListId": 2,\n'
        '    "contactName": "Lucía Gómez",\n'
        '    "phone": "311-234-5678",\n'
        '    "address": "Cra 10 # 5-23",\n'
        '    "initials": "PC",\n'
        '    "active": true,\n'
        '    "createdAt": "2024-01-15",\n'
        '    "branchId": null               // solo asesores\n'
        '  }\n'
        ']'
    )
)

h2('5.2 Obtener usuario por ID')
endpoint_block(
    method='GET',
    path='/users/:id',
    description='Retorna un usuario específico.',
    auth_roles='admin · client/supervisor (solo de su empresa) · propio usuario',
    path_params=':id — ID entero del usuario',
    response_ok='{ "id": 3, "name": "...", "role": "client", ... }',
    response_err='404 { "error": "Usuario no encontrado" }'
)

h2('5.3 Crear usuario')
endpoint_block(
    method='POST',
    path='/users',
    description=(
        'Crea un nuevo usuario. El admin puede crear cualquier rol. '
        'El supervisor solo puede crear usuarios de tipo client en su empresa.'
    ),
    auth_roles='admin · client/supervisor (limitado a role=client, misma empresa)',
    request_body=(
        '{\n'
        '  "name":        "Nombre Apellido",   // requerido\n'
        '  "email":       "user@email.com",    // requerido, único\n'
        '  "password":    "contraseña",        // requerido, min 6 chars\n'
        '  "role":        "client",            // requerido: admin|advisor|client\n'
        '  "clientRole":  "creador_pedidos",   // requerido si role=client: supervisor|creador_pedidos\n'
        '  "companyId":   1,                   // requerido si role=client\n'
        '  "sucursalId":  1,                   // requerido si role=client\n'
        '  "priceListId": 2,                   // requerido si role=client\n'
        '  "branchId":    1,                   // requerido si role=advisor (sede de Papelería Cartagena)\n'
        '  "contactName": "Nombre contacto",   // opcional\n'
        '  "phone":       "311-000-0000",      // opcional\n'
        '  "address":     "Dirección",         // opcional\n'
        '  "active":      true                 // default: true\n'
        '}'
    ),
    response_ok='201 { "id": 5, "name": "...", "email": "...", "role": "client", ... }',
    response_err=(
        '409 { "error": "El email ya está registrado" }\n'
        '403 { "error": "No autorizado para crear este tipo de usuario" }'
    )
)

h2('5.4 Editar usuario')
endpoint_block(
    method='PUT',
    path='/users/:id',
    description='Actualiza datos de un usuario. Si se envía password, debe hashearse antes de guardar.',
    auth_roles='admin · client/supervisor (usuarios de su empresa) · propio usuario (solo sus propios datos básicos)',
    path_params=':id — ID entero del usuario',
    request_body=(
        '{\n'
        '  "name":        "Nuevo nombre",\n'
        '  "email":       "nuevo@email.com",\n'
        '  "password":    "nueva_clave",       // opcional, si se envía se actualiza\n'
        '  "clientRole":  "supervisor",\n'
        '  "sucursalId":  2,\n'
        '  "priceListId": 1,\n'
        '  "branchId":    2,\n'
        '  "active":      false\n'
        '}'
    ),
    response_ok='{ "id": 3, "name": "Nuevo nombre", ... }'
)

h2('5.5 Eliminar usuario')
endpoint_block(
    method='DELETE',
    path='/users/:id',
    description='Desactiva o elimina un usuario. Se recomienda soft-delete.',
    auth_roles='admin · client/supervisor (usuarios de su empresa, no a sí mismo)',
    path_params=':id — ID entero del usuario',
    response_ok='{ "message": "Usuario eliminado correctamente" }',
    response_err='409 { "error": "No se puede eliminar el propio usuario" }'
)

# ── 6. SEDES (PAPELERÍA CARTAGENA) ──────────────────────────────────────────
h1('6. Sedes de Papelería Cartagena')
normal(
    'Las sedes son las instalaciones propias de Papelería Cartagena (no de los clientes). '
    'Se usan para asignar asesores a una sede y para lógistica interna.'
)

h2('6.1 Listar sedes')
endpoint_block(
    method='GET', path='/branches',
    description='Lista todas las sedes de Papelería Cartagena.',
    auth_roles='admin · advisor',
    query_params='active=true|false',
    response_ok=(
        '[\n'
        '  {\n'
        '    "id": 1,\n'
        '    "name": "Sede Centro",\n'
        '    "city": "Bogotá",\n'
        '    "address": "Cra 7 # 15-30",\n'
        '    "phone": "601-234-5678",\n'
        '    "active": true\n'
        '  }\n'
        ']'
    )
)

h2('6.2 Crear sede')
endpoint_block(
    method='POST', path='/branches',
    description='Crea una nueva sede de Papelería Cartagena.',
    auth_roles='admin',
    request_body=(
        '{\n'
        '  "name":    "Sede Sur",       // requerido\n'
        '  "city":    "Bogotá",         // requerido\n'
        '  "address": "Av. 68 # 5-20", // opcional\n'
        '  "phone":   "601-000-0000",  // opcional\n'
        '  "active":  true\n'
        '}'
    ),
    response_ok='201 { "id": 3, ... }'
)

h2('6.3 Editar sede')
endpoint_block(method='PUT', path='/branches/:id', description='Actualiza una sede.', auth_roles='admin',
    path_params=':id — ID entero', request_body='{ "name": "...", "city": "...", "address": "...", "phone": "...", "active": true }',
    response_ok='{ "id": 1, "name": "...", ... }')

h2('6.4 Eliminar sede')
endpoint_block(method='DELETE', path='/branches/:id', description='Elimina o desactiva una sede.', auth_roles='admin',
    path_params=':id', response_ok='{ "message": "Sede eliminada" }',
    response_err='409 { "error": "La sede tiene asesores asignados" }')

# ── 7. CATEGORÍAS ──────────────────────────────────────────────────────────────
h1('7. Categorías de Productos')

h2('7.1 Listar categorías')
endpoint_block(method='GET', path='/categories', description='Lista todas las categorías de productos.',
    auth_roles='admin · advisor · client (todos los autenticados)',
    query_params='active=true|false',
    response_ok='[ { "id": 1, "name": "Papel", "description": "Resmas, pliegos y tipos de papel", "active": true } ]')

h2('7.2 Crear categoría')
endpoint_block(method='POST', path='/categories', description='Crea una nueva categoría.', auth_roles='admin',
    request_body='{\n  "name":        "Cartón",              // requerido\n  "description": "Cajas y embalajes",   // opcional\n  "active":      true\n}',
    response_ok='201 { "id": 6, "name": "Cartón", ... }',
    response_err='409 { "error": "Ya existe una categoría con ese nombre" }')

h2('7.3 Editar categoría')
endpoint_block(method='PUT', path='/categories/:id', description='Actualiza una categoría.', auth_roles='admin',
    path_params=':id — ID entero',
    request_body='{ "name": "...", "description": "...", "active": true }',
    response_ok='{ "id": 1, ... }')

h2('7.4 Eliminar categoría')
endpoint_block(method='DELETE', path='/categories/:id', description='Elimina o desactiva una categoría.', auth_roles='admin',
    path_params=':id',
    response_ok='{ "message": "Categoría eliminada" }',
    response_err='409 { "error": "La categoría tiene productos asociados" }')

# ── 8. LISTAS DE PRECIOS ──────────────────────────────────────────────────────
h1('8. Listas de Precios')
normal(
    'Cada lista tiene un multiplicador que se aplica al precio base del producto. '
    'Los usuarios tipo client tienen un priceListId asignado. El precio final se calcula: '
    'precioFinal = Math.round(basePrice * multiplier).'
)

h2('8.1 Listar listas de precios')
endpoint_block(method='GET', path='/price-lists', description='Lista todas las listas de precios.',
    auth_roles='admin',
    response_ok='[ { "id": 1, "name": "Lista A", "description": "Precio público", "multiplier": 1.0 } ]')

h2('8.2 Crear lista de precios')
endpoint_block(method='POST', path='/price-lists', description='Crea una lista de precios.', auth_roles='admin',
    request_body='{\n  "name":        "Lista D",           // requerido\n  "description": "Precio especial",   // opcional\n  "multiplier":  0.75                 // requerido, 0 < x <= 2\n}',
    response_ok='201 { "id": 4, "name": "Lista D", "multiplier": 0.75 }')

h2('8.3 Editar lista de precios')
endpoint_block(method='PUT', path='/price-lists/:id', description='Actualiza una lista de precios.', auth_roles='admin',
    path_params=':id', request_body='{ "name": "...", "description": "...", "multiplier": 0.85 }',
    response_ok='{ "id": 2, ... }')

h2('8.4 Eliminar lista de precios')
endpoint_block(method='DELETE', path='/price-lists/:id', description='Elimina una lista.', auth_roles='admin',
    path_params=':id', response_ok='{ "message": "Lista eliminada" }',
    response_err='409 { "error": "La lista está asignada a usuarios activos" }')

# ── 9. PRODUCTOS ──────────────────────────────────────────────────────────────
h1('9. Productos (Catálogo)')
normal(
    'Los productos tienen un precio base. El precio mostrado a cada cliente se calcula '
    'según su lista de precios asignada. Los productos complementarios permiten '
    'cross-selling: cuando un cliente ve un producto, se le sugieren los complementarios.'
)

h2('9.1 Listar productos')
endpoint_block(
    method='GET', path='/products',
    description='Lista productos. Si el usuario es client, el precio retornado ya aplica su multiplier.',
    auth_roles='admin · advisor · client (todos los autenticados)',
    query_params=(
        'categoryId=1               — Filtrar por categoría\n'
        'active=true|false          — Filtrar por estado\n'
        'search=texto               — Búsqueda por nombre, SKU o descripción\n'
        'priceListId=2              — (solo admin/advisor) aplicar lista específica al precio\n'
        'page=1&limit=20            — Paginación'
    ),
    response_ok=(
        '{\n'
        '  "data": [\n'
        '    {\n'
        '      "id": 1,\n'
        '      "name": "Resma Papel Bond 75g A4",\n'
        '      "sku": "PAP-001",\n'
        '      "categoryId": 1,\n'
        '      "categoryName": "Papel",              // enriquecido en la respuesta\n'
        '      "description": "Resma de 500 hojas...",\n'
        '      "basePrice": 12500,\n'
        '      "price": 11250,                       // precio con multiplier aplicado\n'
        '      "stock": 150,\n'
        '      "unit": "Resma",\n'
        '      "active": true,\n'
        '      "complementaryIds": [2, 4]            // IDs de productos complementarios\n'
        '    }\n'
        '  ],\n'
        '  "total": 20,\n'
        '  "page": 1,\n'
        '  "limit": 20\n'
        '}'
    )
)

h2('9.2 Obtener producto por ID')
endpoint_block(
    method='GET', path='/products/:id',
    description='Retorna un producto con sus productos complementarios expandidos.',
    auth_roles='admin · advisor · client (todos los autenticados)',
    path_params=':id — ID entero del producto',
    response_ok=(
        '{\n'
        '  "id": 1, "name": "...", "sku": "PAP-001",\n'
        '  "price": 11250,\n'
        '  "complementaries": [\n'
        '    { "id": 2, "name": "Resma Papel Bond 75g Carta", "sku": "PAP-002", "price": 10620 }\n'
        '  ]\n'
        '}'
    ),
    response_err='404 { "error": "Producto no encontrado" }'
)

h2('9.3 Crear producto')
endpoint_block(
    method='POST', path='/products',
    description='Crea un nuevo producto en el catálogo.',
    auth_roles='admin',
    request_body=(
        '{\n'
        '  "name":             "Resma Papel Color",  // requerido\n'
        '  "sku":              "PAP-021",            // requerido, único\n'
        '  "categoryId":       1,                   // requerido\n'
        '  "description":      "Texto libre",        // opcional\n'
        '  "basePrice":        14000,               // requerido, entero positivo\n'
        '  "stock":            100,                 // requerido\n'
        '  "unit":             "Resma",             // requerido (Resma, Caja, Unidad, etc.)\n'
        '  "active":           true,                // default: true\n'
        '  "complementaryIds": [2, 4]              // opcional, array de IDs de productos\n'
        '}'
    ),
    response_ok='201 { "id": 21, "sku": "PAP-021", ... }',
    response_err=(
        '409 { "error": "SKU ya registrado" }\n'
        '404 { "error": "Categoría no encontrada" }'
    )
)

h2('9.4 Editar producto')
endpoint_block(
    method='PUT', path='/products/:id',
    description='Actualiza un producto. Puede usarse también para activar/desactivar (active: false).',
    auth_roles='admin',
    path_params=':id — ID entero',
    request_body=(
        '{\n'
        '  "name":             "Nuevo nombre",\n'
        '  "categoryId":       2,\n'
        '  "description":      "Nueva descripción",\n'
        '  "basePrice":        15000,\n'
        '  "stock":            200,\n'
        '  "unit":             "Resma",\n'
        '  "active":           false,\n'
        '  "complementaryIds": [3, 5, 7]\n'
        '}'
    ),
    response_ok='{ "id": 1, "name": "Nuevo nombre", ... }'
)

h2('9.5 Eliminar producto')
endpoint_block(method='DELETE', path='/products/:id', description='Desactiva (soft-delete) un producto.', auth_roles='admin',
    path_params=':id', response_ok='{ "message": "Producto desactivado" }',
    response_err='409 { "error": "El producto aparece en pedidos activos" }')

# ── 10. PEDIDOS ──────────────────────────────────────────────────────────────
h1('10. Pedidos (Orders)')
normal(
    'Los pedidos son el núcleo del negocio. El flujo de estados es:\n'
    '"Pendiente por aprobar"  →  (supervisor aprueba)  →  "Pendiente"  →\n'
    '"Validar disponibilidad"  →  "Alistamiento"  →  "En Ruta"  →  "Entregado"\n\n'
    'Un pedido creado por creador_pedidos inicia en "Pendiente por aprobar".\n'
    'Un pedido creado por supervisor inicia directo en "Pendiente".\n'
    'Solo los pedidos en estado ≥ "Pendiente" son visibles para el asesor.'
)

h2('10.1 Listar pedidos')
endpoint_block(
    method='GET', path='/orders',
    description=(
        'Lista pedidos con filtros. El admin ve todos. El asesor ve los asignados a él '
        '(status != "Pendiente por aprobar"). El supervisor ve todos los de su empresa. '
        'El creador_pedidos ve solo los suyos propios.'
    ),
    auth_roles='admin · advisor · client/supervisor · client/creador_pedidos (filtrado automático)',
    query_params=(
        'status=Pendiente|En Ruta|...  — Filtrar por estado\n'
        'clientId=3                     — Filtrar por cliente (admin/advisor)\n'
        'advisorId=2                    — Filtrar por asesor (admin)\n'
        'companyId=1                    — Filtrar por empresa (admin)\n'
        'dateFrom=2024-01-01            — Fecha inicio (YYYY-MM-DD)\n'
        'dateTo=2024-12-31              — Fecha fin\n'
        'search=ORD-001                 — Búsqueda por ID\n'
        'page=1&limit=20'
    ),
    response_ok=(
        '{\n'
        '  "data": [\n'
        '    {\n'
        '      "id": "ORD-001",\n'
        '      "clientId": 3,\n'
        '      "clientName": "Papelería El Centro",  // enriquecido\n'
        '      "advisorId": 2,\n'
        '      "advisorName": "Ana Martínez",        // enriquecido\n'
        '      "status": "Entregado",\n'
        '      "createdAt": "2024-02-10",\n'
        '      "updatedAt": "2024-02-12",\n'
        '      "carrier": "TCC",\n'
        '      "total": 150300,\n'
        '      "itemCount": 2,\n'
        '      "commentCount": 2,\n'
        '      "attachmentCount": 1\n'
        '    }\n'
        '  ],\n'
        '  "total": 5, "page": 1, "limit": 20\n'
        '}'
    )
)

h2('10.2 Obtener pedido por ID')
endpoint_block(
    method='GET', path='/orders/:id',
    description='Retorna el pedido completo con items, comentarios y adjuntos.',
    auth_roles='admin · advisor (si está asignado) · client (si es de su empresa)',
    path_params=':id — ID del pedido (ej: "ORD-001")',
    response_ok=(
        '{\n'
        '  "id": "ORD-001",\n'
        '  "clientId": 3, "clientName": "Papelería El Centro",\n'
        '  "advisorId": 2, "advisorName": "Ana Martínez",\n'
        '  "status": "Entregado",\n'
        '  "createdAt": "2024-02-10", "updatedAt": "2024-02-12",\n'
        '  "carrier": "TCC",\n'
        '  "notes": "Entrega urgente solicitada",\n'
        '  "total": 150300,\n'
        '  "items": [\n'
        '    {\n'
        '      "productId": 1,\n'
        '      "sku": "PAP-001",                // enriquecido desde products\n'
        '      "productName": "Resma Papel Bond 75g A4",\n'
        '      "quantity": 10,\n'
        '      "unitPrice": 11250,\n'
        '      "unit": "Resma"\n'
        '    }\n'
        '  ],\n'
        '  "comments": [ { "id": "c1", "authorId": 2, ... } ],\n'
        '  "attachments": [ { "id": "a1", "name": "remision.pdf", ... } ]\n'
        '}'
    ),
    response_err='404 { "error": "Pedido no encontrado" }'
)

h2('10.3 Crear pedido')
endpoint_block(
    method='POST', path='/orders',
    description=(
        'Crea un nuevo pedido. El backend determina automáticamente el status inicial '
        'según el clientRole del usuario autenticado: creador_pedidos → "Pendiente por aprobar", '
        'supervisor → "Pendiente". El advisorId inicial puede ser null.'
    ),
    auth_roles='client/supervisor · client/creador_pedidos',
    request_body=(
        '{\n'
        '  "notes": "Entrega urgente",    // opcional\n'
        '  "items": [                     // requerido, mínimo 1\n'
        '    {\n'
        '      "productId":  1,           // requerido\n'
        '      "quantity":   10,          // requerido, entero positivo\n'
        '      "unitPrice":  11250        // requerido: precio con multiplier ya aplicado\n'
        '    }\n'
        '  ]\n'
        '}'
    ),
    response_ok=(
        '201 {\n'
        '  "id": "ORD-006",\n'
        '  "status": "Pendiente por aprobar",\n'
        '  "total": 112500,\n'
        '  "createdAt": "2024-03-10"\n'
        '}'
    ),
    response_err=(
        '422 { "error": "El pedido debe tener al menos un producto" }\n'
        '404 { "error": "Producto no encontrado: id=99" }'
    ),
    notes_list=(
        'El backend debe generar el ID del pedido (ORD-XXX) con padding a 3 dígitos.\n'
        'El campo "productName" y "unit" de cada item se copian del producto al momento de crear el pedido (snapshot histórico).\n'
        'El total se calcula en el backend: SUM(quantity * unitPrice).'
    )
)

h2('10.4 Actualizar pedido')
endpoint_block(
    method='PUT', path='/orders/:id',
    description=(
        'Actualiza campos de un pedido: estado, transportador, asesor asignado. '
        'El admin puede actualizar todo. El asesor solo puede cambiar status y carrier. '
        'El supervisor puede cambiar status de "Pendiente por aprobar" a "Pendiente" (aprobar) '
        'o a "Rechazado".'
    ),
    auth_roles='admin (todo) · advisor (status, carrier) · client/supervisor (aprobar/rechazar)',
    path_params=':id — ID del pedido',
    request_body=(
        '{\n'
        '  "status":    "En Ruta",      // opcional\n'
        '  "carrier":   "TCC",          // opcional\n'
        '  "advisorId": 2              // opcional, solo admin\n'
        '}'
    ),
    response_ok='{ "id": "ORD-001", "status": "En Ruta", "updatedAt": "2024-03-10", ... }',
    response_err=(
        '403 { "error": "No autorizado para cambiar este estado" }\n'
        '422 { "error": "Transición de estado inválida" }'
    ),
    notes_list=(
        'Validar transiciones de estado permitidas según el rol.\n'
        'Cuando un asesor actualiza un pedido por primera vez, registrar advisorId automáticamente si no estaba asignado.\n'
        'Registrar updatedAt automáticamente en cada actualización.'
    )
)

# ── 11. COMENTARIOS ──────────────────────────────────────────────────────────
h1('11. Comentarios de Pedido')
normal(
    'Los comentarios son internos del equipo (admin + asesores). '
    'Los clientes pueden leerlos pero no crearlos ni eliminarlos. '
    'Son distintos al campo "notes" del pedido, que es una nota libre del cliente.'
)

h2('11.1 Agregar comentario')
endpoint_block(
    method='POST', path='/orders/:orderId/comments',
    description='Agrega un comentario a un pedido.',
    auth_roles='admin · advisor',
    path_params=':orderId — ID del pedido',
    request_body='{\n  "text": "Pedido alistado. Sale en ruta con Envia hoy."  // requerido\n}',
    response_ok=(
        '201 {\n'
        '  "id": "c4",\n'
        '  "orderId": "ORD-001",\n'
        '  "authorId": 2,\n'
        '  "authorName": "Ana Martínez",\n'
        '  "authorRole": "advisor",\n'
        '  "text": "Pedido alistado...",\n'
        '  "createdAt": "2024-03-10T14:30:00"\n'
        '}'
    ),
    response_err='403 { "error": "Los clientes no pueden agregar comentarios" }'
)

h2('11.2 Eliminar comentario')
endpoint_block(
    method='DELETE', path='/orders/:orderId/comments/:commentId',
    description='Elimina un comentario. Solo el autor o el admin pueden eliminarlo.',
    auth_roles='admin · autor del comentario',
    path_params=':orderId, :commentId',
    response_ok='{ "message": "Comentario eliminado" }',
    response_err='403 { "error": "Solo el autor o admin puede eliminar este comentario" }'
)

# ── 12. ADJUNTOS ──────────────────────────────────────────────────────────────
h1('12. Adjuntos de Pedido')
normal(
    'Los adjuntos son archivos subidos por el asesor o admin a un pedido (remisiones, '
    'guías, fotos, etc.). Los clientes pueden verlos y descargarlos pero no subirlos. '
    'Los archivos deben almacenarse en un servicio de storage (S3, GCS, etc.).'
)

warning(
    'Validar tipo y tamaño de archivo en el backend. Tipos permitidos: PDF, JPEG, PNG, XLSX, DOCX. '
    'Tamaño máximo sugerido: 10 MB por archivo.'
)

h2('12.1 Subir adjunto')
endpoint_block(
    method='POST', path='/orders/:orderId/attachments',
    description='Sube un archivo adjunto a un pedido. Usar multipart/form-data.',
    auth_roles='admin · advisor',
    path_params=':orderId — ID del pedido',
    request_body=(
        'Content-Type: multipart/form-data\n'
        '  file: <binario del archivo>   // requerido\n\n'
        'El backend extrae: name, size, type (MIME) del archivo subido.'
    ),
    response_ok=(
        '201 {\n'
        '  "id": "a2",\n'
        '  "orderId": "ORD-001",\n'
        '  "name": "remision-ORD-001.pdf",\n'
        '  "size": 142000,\n'
        '  "type": "application/pdf",\n'
        '  "url": "https://storage.../remision-ORD-001.pdf",  // URL de descarga\n'
        '  "uploadedBy": "Ana Martínez",\n'
        '  "uploadedAt": "2024-03-10T08:20:00"\n'
        '}'
    ),
    response_err=(
        '413 { "error": "Archivo demasiado grande. Máximo 10 MB" }\n'
        '415 { "error": "Tipo de archivo no permitido" }'
    )
)

h2('12.2 Descargar / obtener URL de adjunto')
endpoint_block(
    method='GET', path='/orders/:orderId/attachments/:attachmentId/download',
    description='Retorna una URL de descarga (presigned URL si se usa S3) o redirige al archivo.',
    auth_roles='admin · advisor · client (solo pedidos de su empresa)',
    path_params=':orderId, :attachmentId',
    response_ok=(
        '{ "url": "https://storage.example.com/..?token=xyz&expires=1712000000" }\n'
        '// O redirección 302 directamente al archivo'
    )
)

h2('12.3 Eliminar adjunto')
endpoint_block(
    method='DELETE', path='/orders/:orderId/attachments/:attachmentId',
    description='Elimina un adjunto del pedido y del storage.',
    auth_roles='admin · advisor (solo los que subió)',
    path_params=':orderId, :attachmentId',
    response_ok='{ "message": "Adjunto eliminado" }',
    response_err='403 { "error": "Solo el autor o admin puede eliminar este adjunto" }'
)

# ── 13. CATÁLOGO CON PRECIO ──────────────────────────────────────────────────
h1('13. Catálogo con Precio Aplicado (Cliente)')
normal(
    'Endpoint especializado para el cliente al cargar el catálogo. '
    'Equivale a GET /products pero el precio ya viene calculado con el multiplier '
    'de la lista de precios del usuario autenticado. El frontend no necesita hacer '
    'el cálculo del precio; lo recibe listo para mostrar.'
)

endpoint_block(
    method='GET', path='/catalog',
    description=(
        'Retorna el catálogo de productos activos con el precio ya aplicado según '
        'la lista de precios del usuario autenticado. Incluye los complementarios.'
    ),
    auth_roles='client (supervisor o creador_pedidos)',
    query_params=(
        'categoryId=1    — Filtrar por categoría\n'
        'search=texto    — Búsqueda por nombre o SKU\n'
        'page=1&limit=20'
    ),
    response_ok=(
        '{\n'
        '  "priceListId": 2,\n'
        '  "priceListName": "Lista B",\n'
        '  "data": [\n'
        '    {\n'
        '      "id": 1,\n'
        '      "name": "Resma Papel Bond 75g A4",\n'
        '      "sku": "PAP-001",\n'
        '      "categoryId": 1, "categoryName": "Papel",\n'
        '      "description": "...",\n'
        '      "price": 11250,             // basePrice * multiplier, redondeado\n'
        '      "stock": 150,\n'
        '      "unit": "Resma",\n'
        '      "complementaryIds": [2, 4]\n'
        '    }\n'
        '  ],\n'
        '  "total": 20, "page": 1\n'
        '}'
    ),
    notes_list=(
        'El backend obtiene priceListId del usuario autenticado (JWT).\n'
        'Solo retorna productos con active=true.\n'
        'No exponer basePrice ni multiplier en este endpoint.'
    )
)

# ── 14. ESTADÍSTICAS ──────────────────────────────────────────────────────────
h1('14. Estadísticas y Dashboard')

h2('14.1 Dashboard admin')
endpoint_block(
    method='GET', path='/stats/admin',
    description='Retorna métricas generales para el panel del administrador.',
    auth_roles='admin',
    response_ok=(
        '{\n'
        '  "totalOrders": 5,\n'
        '  "ordersThisMonth": 3,\n'
        '  "totalRevenue": 956100,\n'
        '  "revenueThisMonth": 481500,\n'
        '  "pendingOrders": 2,\n'
        '  "activeClients": 4,\n'
        '  "activeProducts": 20,\n'
        '  "ordersByStatus": {\n'
        '    "Pendiente": 1, "En Ruta": 1, "Entregado": 1, "Alistamiento": 1\n'
        '  },\n'
        '  "revenueByMonth": [\n'
        '    { "month": "2024-01", "revenue": 0 },\n'
        '    { "month": "2024-02", "revenue": 434700 }\n'
        '  ]\n'
        '}'
    )
)

h2('14.2 Dashboard asesor')
endpoint_block(
    method='GET', path='/stats/advisor',
    description='Retorna métricas del asesor autenticado.',
    auth_roles='advisor',
    response_ok=(
        '{\n'
        '  "myOrders": 4,\n'
        '  "myOrdersThisMonth": 2,\n'
        '  "myRevenue": 956100,\n'
        '  "pendingForMe": 1,\n'
        '  "ordersByStatus": { "Pendiente": 1, "En Ruta": 1, ... }\n'
        '}'
    )
)

# ── 15. CONVENCIONES ──────────────────────────────────────────────────────────
h1('15. Convenciones y Reglas Generales')

convs = [
    ('Formato de respuesta',
     'Todas las respuestas exitosas retornan JSON. '
     'Errores retornan { "error": "mensaje" } con el código HTTP apropiado.'),
    ('Códigos HTTP',
     '200 OK · 201 Created · 204 No Content · 400 Bad Request · '
     '401 Unauthorized · 403 Forbidden · 404 Not Found · '
     '409 Conflict · 413 Payload Too Large · 415 Unsupported Media Type · '
     '422 Unprocessable Entity · 500 Internal Server Error'),
    ('Paginación',
     'Todos los endpoints de listado soportan page (default: 1) y limit (default: 20, máx: 100). '
     'Retornar { data: [], total: N, page: N, limit: N }.'),
    ('Fechas',
     'Fechas de solo fecha: YYYY-MM-DD. Fechas con hora: YYYY-MM-DDTHH:MM:SS (ISO 8601). '
     'Siempre en UTC en la base de datos; convertir a zona horaria del cliente en el frontend si es necesario.'),
    ('IDs de pedidos',
     'El ID de pedido sigue el formato ORD-XXX (3 dígitos con cero a la izquierda, ej: ORD-001). '
     'El backend debe generarlos secuencialmente.'),
    ('Soft delete',
     'Para empresas, usuarios, productos y categorías se recomienda soft-delete (campo active=false o deletedAt). '
     'Esto preserva la integridad referencial del historial de pedidos.'),
    ('CORS',
     'Habilitar CORS para el dominio del frontend. En desarrollo: http://localhost:5173 y http://localhost:5174.'),
    ('Variables de entorno',
     'No hardcodear secretos. Usar variables de entorno para: JWT_SECRET, DB_URL, STORAGE_BUCKET, STORAGE_KEY.'),
    ('Autenticación JWT',
     'Expiración sugerida: 8 horas. Incluir refresh token con expiración de 30 días. '
     'Payload mínimo del JWT: { sub: userId, role, clientRole, companyId, exp }.'),
    ('Seguridad de contraseñas',
     'Hashear con bcrypt, mínimo 12 rondas. Nunca almacenar ni retornar contraseñas en texto plano.'),
    ('Validación',
     'Validar todos los inputs en el backend (nunca confiar solo en el frontend). '
     'Retornar errores descriptivos con 422 si hay campos inválidos.'),
    ('Snapshots en pedidos',
     'Al crear un pedido, copiar productName, unit y unitPrice en la tabla order_items. '
     'Estos campos NO deben actualizar si el producto cambia después: son un snapshot histórico.'),
]

tbl = doc.add_table(rows=len(convs)+1, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl.columns[0].width = Cm(5.0)
tbl.columns[1].width = Cm(10.5)
set_cell_bg(tbl.cell(0,0), C_DARK_BLUE); cell_text(tbl.cell(0,0),'Tema',bold=True,color=C_WHITE,size=10)
set_cell_bg(tbl.cell(0,1), C_DARK_BLUE); cell_text(tbl.cell(0,1),'Regla / Convención',bold=True,color=C_WHITE,size=10)
for i,(t,v) in enumerate(convs,1):
    bg = C_LIGHT_BLUE if i % 2 == 0 else C_WHITE
    set_cell_bg(tbl.cell(i,0), bg)
    set_cell_borders(tbl.cell(i,0),'CCCCCC')
    set_cell_borders(tbl.cell(i,1),'CCCCCC')
    cell_text(tbl.cell(i,0), t, bold=True, color=C_MID_BLUE, size=9)
    cell_text(tbl.cell(i,1), v, size=9)

doc.add_paragraph()

# ── RESUMEN ──────────────────────────────────────────────────────────────────
h1('Resumen de Endpoints')

summary = [
    ('POST',   '/auth/login',                               'Público'),
    ('POST',   '/auth/logout',                              'Autenticado'),
    ('GET',    '/auth/me',                                  'Autenticado'),
    ('GET',    '/companies',                                'admin · supervisor'),
    ('POST',   '/companies',                                'admin'),
    ('GET',    '/companies/:id',                            'admin · supervisor'),
    ('PUT',    '/companies/:id',                            'admin'),
    ('DELETE', '/companies/:id',                            'admin'),
    ('GET',    '/companies/:cId/sucursales',                'admin · supervisor'),
    ('POST',   '/companies/:cId/sucursales',                'admin · supervisor'),
    ('PUT',    '/companies/:cId/sucursales/:sId',           'admin · supervisor'),
    ('DELETE', '/companies/:cId/sucursales/:sId',           'admin · supervisor'),
    ('GET',    '/users',                                    'admin · supervisor'),
    ('GET',    '/users/:id',                                'admin · supervisor · propio'),
    ('POST',   '/users',                                    'admin · supervisor'),
    ('PUT',    '/users/:id',                                'admin · supervisor'),
    ('DELETE', '/users/:id',                                'admin · supervisor'),
    ('GET',    '/branches',                                 'admin · advisor'),
    ('POST',   '/branches',                                 'admin'),
    ('PUT',    '/branches/:id',                             'admin'),
    ('DELETE', '/branches/:id',                             'admin'),
    ('GET',    '/categories',                               'todos'),
    ('POST',   '/categories',                               'admin'),
    ('PUT',    '/categories/:id',                           'admin'),
    ('DELETE', '/categories/:id',                           'admin'),
    ('GET',    '/price-lists',                              'admin'),
    ('POST',   '/price-lists',                              'admin'),
    ('PUT',    '/price-lists/:id',                          'admin'),
    ('DELETE', '/price-lists/:id',                          'admin'),
    ('GET',    '/products',                                 'todos'),
    ('GET',    '/products/:id',                             'todos'),
    ('POST',   '/products',                                 'admin'),
    ('PUT',    '/products/:id',                             'admin'),
    ('DELETE', '/products/:id',                             'admin'),
    ('GET',    '/orders',                                   'admin · advisor · client'),
    ('GET',    '/orders/:id',                               'admin · advisor · client'),
    ('POST',   '/orders',                                   'client'),
    ('PUT',    '/orders/:id',                               'admin · advisor · supervisor'),
    ('POST',   '/orders/:oId/comments',                     'admin · advisor'),
    ('DELETE', '/orders/:oId/comments/:cId',                'admin · autor'),
    ('POST',   '/orders/:oId/attachments',                  'admin · advisor'),
    ('GET',    '/orders/:oId/attachments/:aId/download',    'admin · advisor · client'),
    ('DELETE', '/orders/:oId/attachments/:aId',             'admin · advisor'),
    ('GET',    '/catalog',                                  'client'),
    ('GET',    '/stats/admin',                              'admin'),
    ('GET',    '/stats/advisor',                            'advisor'),
]

tbl2 = doc.add_table(rows=len(summary)+1, cols=3)
tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl2.columns[0].width = Cm(1.8)
tbl2.columns[1].width = Cm(8.5)
tbl2.columns[2].width = Cm(5.2)

headers = ['Método','Ruta','Autorización']
for j, h_txt in enumerate(headers):
    set_cell_bg(tbl2.cell(0,j), C_DARK_BLUE)
    cell_text(tbl2.cell(0,j), h_txt, bold=True, color=C_WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

for i, (mth, path, auth) in enumerate(summary, 1):
    bg, fg = METHOD_COLORS.get(mth, (C_GRAY_L, C_GRAY_H))
    row_bg = C_GRAY_L if i % 2 == 0 else C_WHITE
    mth_cell = tbl2.cell(i, 0)
    set_cell_bg(mth_cell, bg)
    set_cell_borders(mth_cell, 'CCCCCC')
    cell_text(mth_cell, mth, bold=True, color=fg, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    for j, val in [(1, path), (2, auth)]:
        c = tbl2.cell(i, j)
        set_cell_bg(c, row_bg)
        set_cell_borders(c, 'CCCCCC')
        cell_text(c, val, size=8.5)

doc.add_paragraph()
normal(f'Total de endpoints: {len(summary)}', bold=True, color=C_MID_BLUE)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = '/Users/nestorhez/Documents/My Code/cartagena/Endpoints_PapeleriaCartagena.docx'
doc.save(output_path)
print(f'✅  Documento guardado en: {output_path}')
